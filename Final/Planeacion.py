import pandas as pd
from docx import Document


# Leer los datos de Excel con pandas
pastYear2 = pd.read_csv('2022.csv', sep=',')
nextYear=pd.read_csv('2023.csv', sep=',')
pastYear = pastYear2.fillna(0) 




# Crear un nuevo documento de Word
doc = Document()
peonesTotal= 0
name=0
yuca=0
platano=0
maiz=0
guandu=0
frijoles=0
arroz=0
#i=0
for i in range(0,71):
    dineroTotal=0
    areaTotal=0
    # Agregar un encabezado al documento
    doc.add_heading('Datos por beneficiario', 0)


    # Agregar una tabla con los datos personales
    table1 = doc.add_table(rows=2, cols=4)
    table1.style = 'Table Grid'


    # Agregar los encabezados de columna a la tabla de datos personales

    table1.cell(0, 0).text = 'Nombre'
    table1.cell(0, 1).text = 'Telefono'
    table1.cell(0, 2).text = 'Area (metros cuadrados)'
    table1.cell(0, 3).text = 'Comunidad'

    #llenado de los datos
    #Tabla 1
    for j in range(0,3):
        table1.cell(1,j).text=str(pastYear.iloc[i,j])




    # Agregar un párrafo con un salto de línea
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()


    # Agregar la segunda tabla al documento con los datos de la cosecha 2022
    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'


    # Agregar los encabezados de columna a la tabla de cosecha 2022
    table2.cell(0, 0).text = 'Rubro sembrado'
    table2.cell(0, 1).text = 'Area sembrada (metros cuadrados)'
    table2.cell(0, 2).text = 'Cantidad cosechada (lb)'
    table2.cell(0, 3).text = 'Rendimiento libra/metro cuadrado'


    # Agregar nombre del rubro
    table2.cell(1, 0).text = 'Maiz'
    table2.cell(2, 0).text = 'Frijoles'
    table2.cell(3, 0).text = 'Aji'
    table2.cell(4, 0).text = 'Platano'
    table2.cell(5, 0).text = 'Yuca'
    table2.cell(6, 0).text = 'Pepino'
    table2.cell(7,0).text='Tomate'


    pastYearNew=pastYear.fillna(0)
    #Llenado de los datos
    #Tabla 2
    #maiz
    table2.cell(1, 1).text = str(pastYearNew.iloc[i,5])
    table2.cell(1, 2).text = str(pastYearNew.iloc[i,6])
    try:
        iloc1 = pastYear.iloc[i,6]
        iloc2 = pastYear.iloc[i,5]
        if(iloc2 == 0):
            table2.cell(1, 3).text = str('0.0')
        else:  
            division = iloc1 / iloc2 
            table2.cell(1, 3).text = str(round(division,4))
    except ZeroDivisionError:
        table2.cell(1, 3).text = str('0.0')
   
    
    #Frijoles
    table2.cell(2, 1).text = str(pastYearNew.iloc[i,11])
    table2.cell(2, 2).text = str(pastYearNew.iloc[i,12])
    try:
        iloc1 = pastYear.iloc[i,12]
        iloc2 = pastYear.iloc[i,11]
        if(iloc2 == 0):
            table2.cell(2, 3).text = str('0.0')
        else:  
            division = iloc1 / iloc2 
            table2.cell(2, 3).text = str(round(division,4))
    except ZeroDivisionError:
        table2.cell(2, 3).text = str('0.0')
    
    
    
    #Aji
    table2.cell(3, 1).text = str(18)
    table2.cell(3, 2).text = str(pastYearNew.iloc[i,18])
    try:
        table2.cell(3, 3).text = str(round(pastYear.iloc[i,18]/18,4))
    except ZeroDivisionError:
        table2.cell(3, 3).text = str('0.0')
    
    
        

    
    #Platano
    table2.cell(4, 1).text = str(pastYearNew.iloc[i,20])
    #table2.cell(1, 2).text = str(pastYear.iloc[i,12])
    #table2.cell(1, 2).text = str(pastYear.iloc[i,12]/pastYear.iloc[i,11])
    #Yuca
    table2.cell(5, 1).text = str(pastYearNew.iloc[i,22])
    #Pepino
    #Area sembrada
    table2.cell(6, 1).text = str(18)
    table2.cell(6, 2).text = str(pastYearNew.iloc[i,14])
    try:
      table2.cell(6,3).text=str(round(pastYear.iloc[i,14]/18,4))
    except ZeroDivisionError:
      table2.cell(6,3).text=str(0.0)
    #rendimiento
    #Tomate
    #Area sembrada
    table2.cell(7, 1).text = str(18)
    table2.cell(7, 2).text = str(pastYearNew.iloc[i,16])
    try:
      table2.cell(7,3).text=str(round(pastYear.iloc[i,16]/18,4))
    except ZeroDivisionError:
      table2.cell(7,3).text=str(0.0)
    #rendimiento




    # Agregar un párrafo con un salto de línea
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()

    table4=doc.add_table(rows=1, cols=4)
    table4.style= "Table Grid"


    table4.cell(0,0).text="Limpieza inicial"
    table4.cell(0,3).text=str(round(pastYear.iloc[i,2]*0.75*8/10000*15+pastYear.iloc[i,2]*0.25*7/10000*15))
    dineroTotal=dineroTotal +pastYear.iloc[i,2]*0.75*8/10000*15+pastYear.iloc[i,2]*0.25*7/10000*15
    # Agregar la segunda tabla al documento con los datos de Excel
    table3 = doc.add_table(rows=15, cols=4)
    table3.style = 'Table Grid'

    # Agregar los encabezados de columna a la tabla de cosecha 2023
    table3.cell(0, 0).text = 'Rubro'
    table3.cell(0, 1).text = 'Area (metros cuadrados)'
    table3.cell(0, 2).text = 'Semilla'
    table3.cell(0, 3).text = 'Pago'
    table3.cell(0, 4).text = 'Total'
    


    # Agregar nombre del rubro
    table3.cell(1, 0).text = 'Pepino'
    table3.cell(2, 0).text = 'Habichuela'
    table3.cell(3, 0).text = 'Aji'
    table3.cell(4, 0).text = 'Culantro'
    table3.cell(5, 0).text = 'Tomate'
    table3.cell(6, 0).text = 'Curcuma y jengibre'
    table3.cell(7, 0).text = 'name'
    table3.cell(8, 0).text = 'Yuca'
    table3.cell(9, 0).text = 'Platano'
    table3.cell(10, 0).text = 'arroz'
    table3.cell(11,0).text=' Guandu'
    table3.cell(12, 0).text = 'Maiz'
    table3.cell(13, 0).text = 'Frijoles'
    table3.cell(14, 0).text = 'Totales'

    #Agregar los datos de cosecha 2023
    #Area
    # Pasar los datos a un array para copiar al documento
    # Limpiar los valores Nan por 0
    
    nextYearNew=nextYear.fillna(0)
    areas=[]
    for j in range(5,19):
        areas.append(nextYearNew.iloc[i,j])
    
    # Copiar los datos al documento
    for j in range (0,14):
        table3.cell(j+1,1).text=str(areas[j])
    # Semilla
    # Hortalizas
    table3.cell(1,2).text=str('54 semillas')   
    table3.cell(2,2).text=str('54 semillas')   
    table3.cell(3,2).text=str('36 matas ')   
    table3.cell(4,2).text=str(' Voleo')   
    table3.cell(5,2).text=str('18 matas ')   
    table3.cell(6,2).text=str('36 matas ')
    # Raices
    table3.cell(7,2).text=str( str(round((areas[6]*2600)/1300)) +' semillas')
    table3.cell(8,2).text=str(str(round((areas[7]*100)/300)) +' matas')
    
    table3.cell(9,2).text=str(str(round((areas[8]*300)/1200)) +' matas')
    table3.cell(10,2).text=str(str(round((areas[9]*3)/300)) +' libras')
    table3.cell(11,2).text=str(str(round((areas[10]*150)/1400)) +' matas')
    table3.cell(12,2).text=str(str(round((areas[11]*15)/10000)) +' Libras')
    table3.cell(13,2).text=str(str(round((areas[12]*30)/10000)) +' Libras')
    

    
    #Inzumo

    #mano de obra
    table3.cell(1,3).text=str('6.42 dolares')   
    table3.cell(2,3).text=str('6.42 dolares')    
    table3.cell(3,3).text=str('12.84 dolares')   
    table3.cell(4,3).text=str('6.42 dolares')     
    table3.cell(4,3).text=str('6.42 dolares')    
    table3.cell(6,3).text=str('6.42 dolares')  
    # Raices
    table3.cell(7,3).text=str(round((areas[6]*10)/10000*15,2) + round((areas[6]*8)/10000*15,2) )
    table3.cell(8,3).text=str(round((areas[7]*10)/10000*15,2) + round((areas[7]*8)/10000*15,2) )
    table3.cell(9,3).text=str(round((areas[8]*10)/10000*15,2) + round((areas[8]*8)/10000*15,2) )
    table3.cell(10,3).text=str(round((areas[9]*7)/10000*15,2) + round((areas[9]*8)/10000*15,2) )
    table3.cell(11,3).text=str(round((areas[10]*7)/10000*15,2) + round((areas[10]*8)/10000*15,2) )
    table3.cell(12,3).text=str(round((areas[11]*7)/10000*15,2) + round((areas[11]*8)/10000*15,2) )
    table3.cell(13,3).text=str(round((areas[12]*7)/10000*15,2) + round((areas[12]*8)/10000*15,2) )
    #Total
    dineroTotal=dineroTotal+45+round((areas[6]*10)/10000*15,2) + round((areas[6]*8)/10000*15,2) + round((areas[7]*10)/10000*15,2) + round((areas[7]*8)/10000*15,2) + round((areas[8]*10)/10000*15,2) + round((areas[8]*8)/10000*15,2) + round((areas[9]*7)/10000*15,2) + round((areas[9]*8)/10000*15,2) + round((areas[10]*7)/10000*15,2) + round((areas[10]*8)/10000*15,2) + round((areas[11]*7)/10000*15,2) + round((areas[11]*8)/10000*15,2) +round((areas[12]*7)/10000*15,2) + round((areas[12]*8)/10000*15,2) 
    areaTotal=  sum(areas)
    table3.cell(14,0).text=str("Totales")
    table3.cell(14,1).text=str(areaTotal)
    table3.cell(14,3).text=str(dineroTotal)

    name = name + round((areas[6]*2600)/1300)
    yuca = yuca + round((areas[7]*100)/300)
    platano = platano + round((areas[8]*300)/1200)
    arroz = arroz + round((areas[9]*3)/300)
    guandu= guandu + round((areas[10]*150)/1400)
    maiz= maiz + round((areas[11]*15)/10000)
    frijoles = frijoles + round((areas[12]*30)/10000)
    peonesTotal=peonesTotal+ dineroTotal

    # Añadir un salto de página después de cada tabl3
    doc.add_paragraph('Beneficiario:_______')
    doc.add_page_break()



# Guardar el documento de Word
doc.save('Datos Beneficiarios.docx')



print ('Resumen ')
print ( 'Pago total: '+ str(peonesTotal))
print ( ' semillas de pepino:'+ str(54*71))
print ( ' semillas de habichuela:'+ str(54*71))
print ( ' matas de aji:'+ str(36*71))
print( ' matas de tomate:'+ str(18*71))
print(' matas de gengibre: '+ str(36*71))
print ( 'semillas de name: '+ str(name)) 
print( 'matas de yuca: '+ str(yuca))
print('matas de platano: )'+ str(platano))
print ('libras de arroz: ' +str(arroz))
print (' matas de guandu: '+ str(guandu))
print(' libras de maiz:' + str(maiz))

print('libras de frijoles'+ str(frijoles))