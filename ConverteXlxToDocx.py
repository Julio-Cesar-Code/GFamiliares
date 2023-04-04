import pandas as pd
from docx import Document
from docx.shared import Inches

# Leer los datos de Excel con pandas
pastYear = pd.read_csv('Past.csv', sep=',')
nextYear=pd.read_csv('Programming.csv', sep=',')





# Crear un nuevo documento de Word
doc = Document()


#i=0
for i in range(0,30):
    # Agregar un encabezado al documento
    doc.add_heading('Datos por beneficiario', 0)


    # Agregar una tabla con los datos personales
    table1 = doc.add_table(rows=2, cols=3)
    table1.style = 'Table Grid'


    # Agregar los encabezados de columna a la tabla de datos personales

    table1.cell(0, 0).text = 'Nombre'
    table1.cell(0, 1).text = 'Telefono'
    table1.cell(0, 2).text = 'Area (metros cuadrados)'

    #llenado de los datos
    #Tabla 1
    for j in range(0,3):
        table1.cell(1,j).text=str(pastYear.iloc[i,j])




    # Agregar un párrafo con un salto de línea
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()


    # Agregar la segunda tabla al documento con los datos de la cosecha 2022
    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'


    # Agregar los encabezados de columna a la tabla de cosecha 2022
    table2.cell(0, 0).text = 'Rubro sembrado'
    table2.cell(0, 1).text = 'Area sembrada (metros cuadrados)'
    table2.cell(0, 2).text = 'Cantidad cosechada (lb)'
    table2.cell(0, 3).text = 'Rendimiento libra/metro cuadrado'


    # Agregar nombre del rubro
    table2.cell(1, 0).text = 'Maiz'
    table2.cell(2, 0).text = 'Frijoles'
    table2.cell(3, 0).text = 'Aji'
    table2.cell(4, 0).text = 'Platano'
    table2.cell(5, 0).text = 'Yuca'
    table2.cell(6, 0).text = 'Pepino'
    table2.cell(7,0).text='Tomate'


   
    #Llenado de los datos
    #Tabla 2
    #maiz
    table2.cell(1, 1).text = str(pastYear.iloc[i,5])
    table2.cell(1, 2).text = str(pastYear.iloc[i,6])
    try:
        table2.cell(1, 3).text = str(round(pastYear.iloc[i,6]/pastYear.iloc[i,5],4))
    except ValueError:
        print("Esta operacion no se puede realizar")
   
    
    #Frijoles
    table2.cell(2, 1).text = str(pastYear.iloc[i,11])
    table2.cell(2, 2).text = str(pastYear.iloc[i,12])
    try:
        table2.cell(2, 3).text = str(round(pastYear.iloc[i,12]/pastYear.iloc[i,11],4))
    except ValueError:
        print("Esta operacion no se puede realizar")
    
    
    
    #Aji
    table2.cell(3, 1).text = str(18)
    table2.cell(3, 2).text = str(pastYear.iloc[i,18])
    try:
        table2.cell(3, 3).text = str(round(pastYear.iloc[i,18]/18,4))
    except ValueError:
        print("Esta operacion no se puede realizar")
    
    
        

    
    #Platano
    table2.cell(4, 1).text = str(pastYear.iloc[i,20])
    #table2.cell(1, 2).text = str(pastYear.iloc[i,12])
    #table2.cell(1, 2).text = str(pastYear.iloc[i,12]/pastYear.iloc[i,11])
    #Yuca
    table2.cell(5, 1).text = str(pastYear.iloc[i,22])
    #Pepino
    #Area sembrada
    table2.cell(6, 1).text = str(18)
    table2.cell(6, 2).text = str(pastYear.iloc[i,14])
    #rendimiento
    #Tomate
    #Area sembrada
    table2.cell(7, 1).text = str(18)
    table2.cell(7, 2).text = str(pastYear.iloc[i,16])
    #rendimiento




    # Agregar un párrafo con un salto de línea
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()


    # Agregar la segunda tabla al documento con los datos de Excel
    table3 = doc.add_table(rows=15, cols=6)
    table3.style = 'Table Grid'

    # Agregar los encabezados de columna a la tabla de cosecha 2023
    table3.cell(0, 0).text = 'Rubro a sembrar'
    table3.cell(0, 1).text = 'Area (metros cuadrados)'
    table3.cell(0, 2).text = 'Cantidad de semilla'
    table3.cell(0, 3).text = 'Inzumo necesario'
    table3.cell(0, 4).text = 'Mano de obra ($)'
    table3.cell(0, 5).text = 'Cosecha 2023'


    # Agregar nombre del rubro
    table3.cell(1, 0).text = 'Pepino'
    table3.cell(2, 0).text = 'Habichuela'
    table3.cell(3, 0).text = 'Aji'
    table3.cell(4, 0).text = 'Culantro'
    table3.cell(5, 0).text = 'Tomate'
    table3.cell(6, 0).text = 'Curcuma y jengibre'
    table3.cell(7, 0).text = 'name'
    table3.cell(8, 0).text = 'Yuca'
    table3.cell(9, 0).text = 'Platano'
    table3.cell(10, 0).text = 'nampi'
    table3.cell(11, 0).text = 'arroz'
    table3.cell(12,0).text=' Guandu'
    table3.cell(13, 0).text = 'Maiz'
    table3.cell(14, 0).text = 'Frijoles'

    #Agregar los datos de cosecha 2023
    #Area
    # Pasar los datos a un array para copiar al documento
    # Limpiar los valores Nan por 0
    
    nextYearNew=nextYear.fillna(0)
    areas=[]
    for j in range(9,23):
        areas.append(nextYearNew.iloc[i,j])
    
    # Copiar los datos al documento
    for j in range (0,14):
        table3.cell(j+1,1).text=str(areas[j])
    # Semilla
    # Hortalizas
    table3.cell(1,2).text=str('100 semillas')   
    table3.cell(2,2).text=str('75 semillas')   
    table3.cell(3,2).text=str('20 matas ')   
    table3.cell(4,2).text=str(' ')   
    table3.cell(5,2).text=str('20 matas ')   
    table3.cell(6,2).text=str('50 matas ')
    # Raices
    table3.cell(7,2).text=str( str(round((areas[6]*2600)/1300)) +' semillas')
    table3.cell(8,2).text=str(str(round((areas[7]*100)/300)) +' matas')
    
    table3.cell(9,2).text=str(str(round((areas[8]*300)/1200)) +' matas')
    table3.cell(11,2).text=str(str(round((areas[10]*3)/300)) +' libras')
    table3.cell(12,2).text=str(str(round((areas[11]*150)/1400)) +' matas')
    table3.cell(13,2).text=str(str(round((areas[12]*15)/10000)) +' Libras')
    table3.cell(14,2).text=str(str(round((areas[13]*30)/10000)) +' Libras')
    

    
    #Inzumo

    #mano de obra
    table3.cell(1,4).text=str('7.5 dolares')   
    table3.cell(2,4).text=str('7.5 dolares')    
    table3.cell(3,4).text=str('7.5 dolares')   
    table3.cell(4,4).text=str('7.5 dolares')     
    table3.cell(4,4).text=str('7.5 dolares')    
    table3.cell(6,4).text=str('7.5 dolares')  
    # Raices
    table3.cell(7,4).text=str(' M=' + str(round((areas[6]*10)/10000*15,2)) +', s= '+ str(round((areas[6]*10)/10000*15,2)) )
    table3.cell(8,4).text=str(' M=' + str(round((areas[7]*10)/10000*15,2)) +', s= '+ str(round((areas[7]*10)/10000*15,2)) )
    table3.cell(9,4).text=str(' M=' + str(round((areas[8]*10)/10000*15,2)) +', s= '+ str(round((areas[8]*10)/10000*15,2)) )
    table3.cell(11,4).text=str(' M=' + str(round((areas[10]*8)/10000*15,2)) +', s= '+ str(round((areas[10]*7)/10000*15,2)) )
    table3.cell(12,4).text=str(' M=' + str(round((areas[11]*8)/10000*15,2)) +', s= '+ str(round((areas[11]*7)/10000*15,2)) )
    table3.cell(13,4).text=str(' M=' + str(round((areas[12]*8)/10000*15,2)) +', s= '+ str(round((areas[12]*7)/10000*15,2)) )
    table3.cell(14,4).text=str(' M=' + str(round((areas[13]*8)/10000*15,2)) +', s= '+ str(round((areas[13]*7)/10000*15,2)) )
    #Cosecha


    # Añadir un salto de página después de cada tabla
    doc.add_page_break()



# Guardar el documento de Word
doc.save('Datos Beneficiarios.docx')